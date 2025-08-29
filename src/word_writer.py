from __future__ import annotations

import logging
import re
from typing import Dict

from docx import Document


class WordWriter:
    """
    Substituição rápida via regex unificado (evita loops O(n*m)).
    Por quê: performance consistente com muitos marcadores.
    """

    def __init__(self, template_path):
        self.template_path = str(template_path)
        self.logger = logging.getLogger(__name__)

    def replace_in_document(self, replacements: Dict[str, str], output_path) -> bool:
        try:
            doc = Document(self.template_path)
        except Exception as e:
            self.logger.error(f"Erro ao abrir o modelo Word: {e}")
            return False

        try:
            pattern = self._build_union_pattern(replacements)
            self._replace_in_paragraphs(doc.paragraphs, pattern, replacements)
            for table in doc.tables:
                self._replace_in_table(table, pattern, replacements)
            for section in doc.sections:
                self._replace_in_paragraphs(section.header.paragraphs, pattern, replacements)
                self._replace_in_paragraphs(section.footer.paragraphs, pattern, replacements)
                for t in section.header.tables:
                    self._replace_in_table(t, pattern, replacements)
                for t in section.footer.tables:
                    self._replace_in_table(t, pattern, replacements)
            doc.save(str(output_path))
            self.logger.info(f"Contrato salvo em: {output_path}")
            return True
        except Exception as e:
            self.logger.error(f"Falha na geração do documento: {e}")
            return False

    @staticmethod
    def _build_union_pattern(replacements: Dict[str, str]) -> re.Pattern[str]:
        keys = sorted(replacements.keys(), key=len, reverse=True)
        union = "|".join(re.escape(k) for k in keys)
        return re.compile(union)

    def _replace_text(self, text: str, pattern, replacements: Dict[str, str]) -> str:
        return pattern.sub(lambda m: replacements.get(m.group(0), ""), text)

    def _replace_in_table(self, table, pattern, replacements: Dict[str, str]) -> None:
        for row in table.rows:
            for cell in row.cells:
                self._replace_in_paragraphs(cell.paragraphs, pattern, replacements)
                for inner in cell.tables:
                    self._replace_in_table(inner, pattern, replacements)

    def _replace_in_paragraphs(self, paragraphs, pattern, replacements: Dict[str, str]) -> None:
        for p in paragraphs:
            new_text = self._replace_text(p.text, pattern, replacements)
            if new_text != p.text:
                p.text = new_text  # reset de runs é intencional